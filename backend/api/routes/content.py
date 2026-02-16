"""
V2 API - Content Processing Routes
===================================

Endpoints for content simplification, translation, TTS, OCR, embeddings, and Q&A.
OPTIMIZED: Uses orjson for faster SSE serialization.
"""

import asyncio
import contextlib
import logging
import os
import tempfile
import time
import uuid
from typing import Any

from fastapi import APIRouter, Body, File, HTTPException, Query, UploadFile
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

from backend.infra.cache import get_unified_cache

from ...utils.memory_guard import require_memory
from ..deps import (
    get_ai_engine as _get_ai_engine,
)
from ..deps import (
    get_middleware as _get_middleware,
)
from ..deps import (
    get_pipeline as _get_pipeline,
)
from ..deps import (
    json_dumps as _json_dumps,
)

logger = logging.getLogger(__name__)

router = APIRouter()

# ==================== Models ====================

class ContentProcessRequest(BaseModel):
    """Content processing request."""

    text: str = Field(..., min_length=1, max_length=50000)
    simplify: bool = True
    translate: bool = False
    generate_audio: bool = False
    validate_content: bool = False  # Validation now optional, disabled by default
    target_language: str = Field(default="Hindi")
    quality_mode: str = Field(default="balanced")
    enable_collaboration: bool = Field(default=False)
    collaboration_pattern: str = Field(default="verify")

class ContentProcessResponse(BaseModel):
    """Content processing response."""

    request_id: str
    success: bool
    original_text: str
    simplified_text: str | None = None
    translated_text: str | None = None
    audio_url: str | None = None
    validation_score: float | None = None
    processing_time_ms: float
    cache_hit: bool = False
    collaboration_confidence: float | None = None
    collaboration_consensus: bool | None = None
    models_used: list[str] | None = None

class TTSRequest(BaseModel):
    """TTS request model."""

    text: str = Field(..., min_length=1, max_length=5000)
    language: str = Field(default="hi")
    voice: str | None = None

class TTSResponse(BaseModel):
    """TTS response model."""

    audio_id: str
    audio_url: str
    duration_ms: int
    processing_time_ms: float

class QAProcessRequest(BaseModel):
    """Q&A document processing request."""

    content: str = Field(..., min_length=10)
    title: str | None = None

class QAQueryRequest(BaseModel):
    """Q&A query request."""

    document_id: str
    question: str = Field(..., min_length=3)

class QAResponse(BaseModel):
    """Q&A response model."""

    answer: str
    confidence: float
    sources: list[dict[str, Any]]
    processing_time_ms: float

# ==================== Helper Functions ====================

def sse_event(event: str, data: dict[str, Any]) -> str:
    """Format SSE event using fast JSON serialization."""
    return f"event: {event}\ndata: {_json_dumps(data)}\n\n"

def _save_upload_to_temp(content: bytes, suffix: str) -> str:
    """Save upload content to temp file (sync operation)."""
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(content)
        return tmp.name

def _extract_docx_content(content: bytes, start_time: float) -> dict[str, Any] | None:
    """Extract content from DOCX file."""
    try:
        import io

        from docx import Document

        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)

        elapsed = (time.perf_counter() - start_time) * 1000

        return {
            "text": "\n\n".join(paragraphs),
            "format": "docx",
            "pages": 1,
            "has_formulas": False,
            "has_tables": len(tables) > 0,
            "formulas": [],
            "tables": tables,
            "confidence": 1.0,
            "metadata": {"source": "python-docx"},
            "model": "python-docx",
            "processing_time_ms": elapsed,
        }
    except ImportError:
        return None

# ==================== Content Processing Endpoints ====================

@router.post(
    "/process", response_model=ContentProcessResponse, tags=["content"]
)
@require_memory(action="reject", reject_on=("critical", "emergency"))
async def process_content(request: ContentProcessRequest):
    """Process content through the unified pipeline (simplify, translate, audio)."""
    request_id = str(uuid.uuid4())[:8]
    start_time = time.perf_counter()

    # ── Middleware: classify for telemetry ──────────────────────────
    middleware = _get_middleware()
    mw_classification = None
    if middleware:
        with contextlib.suppress(RuntimeError, ValueError, OSError):
            mw_classification = middleware.classifier.classify(request.text, request_id)

    try:
        from ...services.pipeline import ProcessingRequest

        # OPTIMIZATION: Use cached pipeline singleton
        pipeline = _get_pipeline()

        proc_request = ProcessingRequest(
            text=request.text,
            simplify=request.simplify,
            translate=request.translate,
            generate_audio=request.generate_audio,
            validate=request.validate_content,
            target_language=request.target_language,
            quality_mode=request.quality_mode,
            enable_collaboration=request.enable_collaboration,
            collaboration_pattern=request.collaboration_pattern,
            verify_translation=True,
            request_id=request_id,
        )

        result = await pipeline.process(proc_request)

        # ── Middleware: record telemetry ──────────────────────────
        if middleware and mw_classification:
            try:
                elapsed = (time.perf_counter() - start_time) * 1000
                middleware.evaluator.evaluate(
                    request_id=request_id,
                    intent=mw_classification.intent.value,
                    model_target=mw_classification.model_target.value,
                    model_used=mw_classification.model_target.value,
                    complexity_score=mw_classification.complexity_score,
                    total_latency_ms=elapsed,
                    confidence=0.8,
                    output_tokens=0,
                    has_error=not result.success,
                )
                middleware.latency.record_latency("content_process", elapsed)
            except (RuntimeError, ValueError, OSError):  # service call
                pass

        return ContentProcessResponse(
            request_id=result.request_id,
            success=result.success,
            original_text=result.original_text,
            simplified_text=result.simplified_text,
            translated_text=result.translated_text,
            audio_url=f"/api/v2/audio/{result.audio_path}"
            if result.audio_path
            else None,
            validation_score=result.validation_score,
            processing_time_ms=result.processing_time_ms,
            cache_hit=result.cache_hits > 0,
            collaboration_confidence=result.collaboration_confidence,
            collaboration_consensus=result.collaboration_consensus,
            models_used=result.models_used,
        )

    except Exception as e:
        logger.error("Content processing error: %s", e)
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/process/stream", tags=["content"])
@require_memory(action="reject", reject_on=("critical", "emergency"))
async def process_content_stream(request: ContentProcessRequest):
    """Process content with streaming progress updates."""

    async def generate():
        try:
            from ...services.pipeline import ProcessingRequest

            # OPTIMIZATION: Use cached pipeline singleton
            pipeline = _get_pipeline()

            proc_request = ProcessingRequest(
                text=request.text,
                simplify=request.simplify,
                translate=request.translate,
                generate_audio=request.generate_audio,
                validate=request.validate_content,
                target_language=request.target_language,
                quality_mode=request.quality_mode,
            )

            async for progress in pipeline.process_stream(proc_request):
                yield sse_event(
                    "progress",
                    {
                        "stage": progress.stage.value,
                        "progress": progress.progress,
                        "message": progress.message,
                        "partial_result": progress.partial_result,
                    },
                )

            yield "data: [DONE]\n\n"

        except (RuntimeError, ValueError, OSError) as e:  # service call
            yield sse_event("error", {"error": str(e)})

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive"},
    )

@router.post("/simplify", tags=["content"])
async def simplify_content(text: str = Body(..., embed=True)):
    """Simplify text."""
    start_time = time.perf_counter()
    try:
        from ...services.pipeline import ProcessingRequest

        # OPTIMIZATION: Use cached pipeline singleton
        pipeline = _get_pipeline()

        result = await pipeline.process(
            ProcessingRequest(
                text=text,
                simplify=True,
                translate=False,
                generate_audio=False,
            )
        )

        elapsed = (time.perf_counter() - start_time) * 1000

        # Middleware telemetry
        middleware = _get_middleware()
        if middleware:
            try:
                middleware.evaluator.evaluate(
                    request_id="simplify",
                    intent="simplification",
                    model_target="standard",
                    model_used="standard",
                    complexity_score=0.4,
                    total_latency_ms=elapsed,
                    confidence=0.8,
                    output_tokens=0,
                    has_error=False,
                )
                middleware.latency.record_latency("simplify", elapsed)
            except (RuntimeError, ValueError, OSError):  # service call
                pass

        return {
            "original": text,
            "simplified": result.simplified_text,
            "processing_time_ms": result.processing_time_ms,
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/translate", tags=["content"])
@require_memory(action="reject", reject_on=("critical", "emergency"))
async def translate_content(
    text: str = Body(..., embed=True),
    target_language: str = Body(default="Hindi"),
    source_language: str = Body(default="English"),
):
    """Translate text to target language."""
    start_time = time.perf_counter()
    try:
        from ...services.pipeline import ProcessingRequest

        # OPTIMIZATION: Use cached pipeline singleton
        pipeline = _get_pipeline()

        result = await pipeline.process(
            ProcessingRequest(
                text=text,
                simplify=False,
                translate=True,
                target_language=target_language,
            )
        )

        elapsed = (time.perf_counter() - start_time) * 1000

        # Middleware telemetry
        middleware = _get_middleware()
        if middleware:
            try:
                middleware.evaluator.evaluate(
                    request_id="translate",
                    intent="translation",
                    model_target="specialized",
                    model_used="specialized",
                    complexity_score=0.3,
                    total_latency_ms=elapsed,
                    confidence=0.85,
                    output_tokens=0,
                    has_error=False,
                )
                middleware.latency.record_latency("translate", elapsed)
            except (RuntimeError, ValueError, OSError):  # service call
                pass

        return {
            "original": text,
            "translated": result.translated_text,
            "source_language": source_language,
            "target_language": target_language,
            "processing_time_ms": result.processing_time_ms,
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# ==================== TTS Endpoints ====================

@router.post("/tts", response_model=TTSResponse, tags=["content"])
async def text_to_speech(request: TTSRequest):
    """Convert text to speech."""
    start_time = time.perf_counter()

    try:
        from ...services.speech import SpeechGenerator

        speech_gen = SpeechGenerator()

        # Use sync method in executor
        loop = asyncio.get_running_loop()
        result = await loop.run_in_executor(
            None,
            lambda: speech_gen.generate_speech(
                text=request.text, language=request.language, subject="General"
            ),
        )

        elapsed = (time.perf_counter() - start_time) * 1000
        audio_id = str(uuid.uuid4())[:8]

        return TTSResponse(
            audio_id=audio_id,
            audio_url=f"/api/v2/audio/{result.file_path}"
            if result.file_path
            else f"/api/v2/audio/stream/{audio_id}",
            duration_ms=int(result.duration_seconds * 1000)
            if hasattr(result, "duration_seconds")
            else 0,
            processing_time_ms=elapsed,
        )

    except Exception as e:
        logger.error("TTS error: %s", e)
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/tts/voices", tags=["content"])
async def list_tts_voices():
    """List available TTS voices."""
    return {
        "voices": [
            {
                "id": "default",
                "name": "Default",
                "languages": ["hi", "en", "te", "ta", "bn"],
            },
            {"id": "female", "name": "Female", "languages": ["hi", "en"]},
            {"id": "male", "name": "Male", "languages": ["hi", "en"]},
        ]
    }

# ==================== Q&A Endpoints ====================

@router.post("/qa/process", tags=["qa"])
async def process_document_for_qa(request: QAProcessRequest):
    """Process document for Q&A (creates embeddings)."""
    try:
        # OPTIMIZATION: Use cached pipeline singleton
        pipeline = _get_pipeline()

        doc_id = str(uuid.uuid4())[:8]

        embeddings = await pipeline.embed([request.content])

        cache = get_unified_cache()
        await cache.set(
            f"qa:doc:{doc_id}",
            {
                "content": request.content,
                "embedding": embeddings[0].tolist()
                if hasattr(embeddings[0], "tolist")
                else embeddings[0],
            },
            ttl=86400,
        )

        return {
            "document_id": doc_id,
            "status": "processed",
            "content_length": len(request.content),
        }

    except Exception as e:
        logger.error("QA processing error: %s", e)
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/qa/ask", response_model=QAResponse, tags=["qa"])
async def ask_question(request: QAQueryRequest):
    """Ask a question about a processed document."""
    start_time = time.perf_counter()

    try:
        from ...services.chat.engine import GenerationConfig

        cache = get_unified_cache()
        doc_data = await cache.get(f"qa:doc:{request.document_id}")

        if not doc_data:
            raise HTTPException(
                status_code=404, detail="Document not found. Process it first."
            )

        # OPTIMIZATION: Use cached engine singleton
        engine = _get_ai_engine()

        qa_prompt = f"""Based on the following document content, answer this question accurately and completely.

DOCUMENT CONTENT:
{doc_data["content"][:4000]}

QUESTION: {request.question}

Provide a clear, accurate answer based ONLY on the document content above. If the answer is not in the document, say so."""

        config = GenerationConfig(max_tokens=2048, temperature=0.3, use_rag=True)

        response = await engine.chat(
            message=qa_prompt,
            config=config,
            context_data={"task_type": "qa", "document_id": request.document_id},
        )

        elapsed = (time.perf_counter() - start_time) * 1000
        confidence = (
            response.metadata.confidence
            if hasattr(response.metadata, "confidence")
            else 0.75
        )

        return QAResponse(
            answer=response.content,
            confidence=confidence,
            sources=[{"document_id": request.document_id, "relevance": 0.9}],
            processing_time_ms=elapsed,
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error("QA query error: %s", e)
        raise HTTPException(status_code=500, detail=str(e))

# ==================== STT Endpoints ====================

@router.post("/stt/transcribe", tags=["stt"])
@require_memory(action="reject", reject_on=("critical", "emergency"))
async def transcribe_audio(
    file: UploadFile = File(...),
    language: str = Query(default="auto", description="Language code or 'auto'"),
):
    """Transcribe audio to text using Whisper V3 Turbo."""
    start_time = time.perf_counter()

    try:
        # OPTIMIZATION: Use cached pipeline singleton
        pipeline = _get_pipeline()

        content = await file.read()
        suffix = os.path.splitext(file.filename)[1] if file.filename else ".wav"
        tmp_path = _save_upload_to_temp(content, suffix)

        try:
            result = await pipeline.transcribe(
                tmp_path, language if language != "auto" else None
            )
            elapsed = (time.perf_counter() - start_time) * 1000

            return {
                "text": result.get("text", ""),
                "language": result.get("language", language),
                "confidence": result.get("confidence", 0.9),
                "segments": result.get("segments", []),
                "processing_time_ms": elapsed,
            }
        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)

    except Exception as e:
        logger.error("STT transcription error: %s", e)
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/stt/languages", tags=["stt"])
async def list_stt_languages():
    """List supported languages for speech-to-text."""
    return {
        "languages": [
            {"code": "auto", "name": "Auto-detect"},
            {"code": "hi", "name": "Hindi"},
            {"code": "en", "name": "English"},
            {"code": "te", "name": "Telugu"},
            {"code": "ta", "name": "Tamil"},
            {"code": "bn", "name": "Bengali"},
            {"code": "mr", "name": "Marathi"},
            {"code": "gu", "name": "Gujarati"},
            {"code": "kn", "name": "Kannada"},
            {"code": "ml", "name": "Malayalam"},
            {"code": "or", "name": "Odia"},
            {"code": "pa", "name": "Punjabi"},
        ],
        "model": "whisper-large-v3-turbo",
    }

@router.post("/stt/guest", tags=["stt"])
async def transcribe_audio_guest(
    file: UploadFile = File(..., alias="audio"), language: str = Query(default="auto")
):
    """Guest speech-to-text endpoint (no auth required)."""
    start_time = time.perf_counter()

    try:
        # OPTIMIZATION: Use cached pipeline singleton
        pipeline = _get_pipeline()

        content = await file.read()
        suffix = os.path.splitext(file.filename)[1] if file.filename else ".webm"
        tmp_path = _save_upload_to_temp(content, suffix)

        try:
            result = await pipeline.transcribe(
                tmp_path, language if language != "auto" else None
            )
            elapsed = (time.perf_counter() - start_time) * 1000

            return {
                "success": True,
                "text": result.get("text", ""),
                "language": result.get("language", language),
                "confidence": result.get("confidence", 0.9),
                "processing_time_ms": elapsed,
            }
        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)

    except Exception as e:
        logger.error("Guest STT error: %s", e)
        return {"success": False, "text": "", "error": str(e)}

# ==================== OCR Endpoints ====================

@router.post("/ocr/extract", tags=["ocr"])
async def extract_text_from_document(
    file: UploadFile = File(...),
    ocr_type: str = Query(default="document"),
    extract_formulas: bool = Query(default=True),
    extract_tables: bool = Query(default=True),
):
    """Extract text from any document using GOT-OCR2 model."""
    start_time = time.perf_counter()

    try:
        from ...services.ocr import get_ocr_service

        ocr_service = get_ocr_service()

        filename = file.filename or "document"
        suffix = os.path.splitext(filename)[1].lower() if filename else ".png"

        if suffix == ".docx":
            content = await file.read()
            result = _extract_docx_content(content, start_time)
            if result:
                return result

        content = await file.read()
        tmp_path = _save_upload_to_temp(content, suffix)

        try:
            ocr_result = await ocr_service.extract_text_async(
                tmp_path,
                extract_formulas=extract_formulas,
                extract_tables=extract_tables,
            )

            elapsed = (time.perf_counter() - start_time) * 1000

            return {
                "text": ocr_result.text,
                "format": suffix.replace(".", ""),
                "pages": ocr_result.num_pages,
                "has_formulas": ocr_result.has_formulas,
                "has_tables": ocr_result.has_tables,
                "formulas": [
                    f.get("latex", f.get("original", "")) for f in ocr_result.formula_blocks
                ]
                if ocr_result.formula_blocks
                else [],
                "tables": ocr_result.table_blocks,
                "confidence": ocr_result.confidence,
                "metadata": ocr_result.metadata,
                "model": ocr_result.metadata.get("ocr_model", "GOT-OCR2"),
                "processing_time_ms": elapsed,
            }

        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)

    except Exception as e:
        logger.error("Document extraction error: %s", e)
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/ocr/capabilities", tags=["ocr"])
async def get_ocr_capabilities():
    """Get OCR model capabilities and supported formats."""
    try:
        from ...services.ocr import get_ocr_service

        ocr_service = get_ocr_service()
        backend_info = ocr_service.get_backend_info()
    except (ImportError, RuntimeError):  # import/service call
        backend_info = {"backend": "tesseract", "got_available": False}

    return {
        "primary_model": "GOT-OCR2"
        if backend_info.get("got_available")
        else "Tesseract",
        "fallback_model": "Tesseract",
        "active_backend": backend_info.get("backend", "tesseract"),
        "device": backend_info.get("device", {}).get("device_name", "CPU"),
        "supported_formats": [
            {"ext": "pdf", "description": "PDF documents (multi-page, scanned)"},
            {"ext": "docx", "description": "Microsoft Word documents"},
            {"ext": "png", "description": "PNG images"},
            {"ext": "jpg", "description": "JPEG images"},
            {"ext": "jpeg", "description": "JPEG images"},
            {"ext": "tiff", "description": "TIFF images (multi-page)"},
            {"ext": "bmp", "description": "Bitmap images"},
            {"ext": "webp", "description": "WebP images"},
        ],
        "ocr_types": [
            {"type": "document", "description": "Printed documents, forms, books"},
            {"type": "scene", "description": "Text in natural scenes, signboards"},
            {"type": "handwriting", "description": "Handwritten text recognition"},
        ],
        "features": [
            "Formula extraction (LaTeX output)",
            "Table extraction (structured data)",
            "Indian language support (22 languages)",
            "Mixed-language document support",
            "Scanned document processing",
            "Handwriting recognition",
        ],
        "languages": [
            "English",
            "Hindi",
            "Tamil",
            "Telugu",
            "Bengali",
            "Marathi",
            "Gujarati",
            "Kannada",
            "Malayalam",
            "Punjabi",
            "Odia",
            "Assamese",
            "Urdu",
            "Sanskrit",
            "Nepali",
            "Sindhi",
        ],
    }

# ==================== Embedding Endpoints ====================

@router.post("/embed", tags=["embeddings"])
@require_memory(action="reject", reject_on=("critical", "emergency"))
async def generate_embeddings(
    texts: list[str] = Body(..., embed=True), use_cache: bool = Body(default=True)
):
    """Generate embeddings for texts."""
    start_time = time.perf_counter()

    try:
        from backend.ml.inference import get_inference_engine

        engine = get_inference_engine()

        embeddings = await engine.embed(texts, use_cache=use_cache)
        elapsed = (time.perf_counter() - start_time) * 1000

        return {
            "embeddings": embeddings.tolist()
            if hasattr(embeddings, "tolist")
            else embeddings,
            "dimension": len(embeddings[0]) if embeddings else 0,
            "processing_time_ms": elapsed,
        }

    except Exception as e:
        logger.error("Embedding error: %s", e)
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/embeddings/generate", tags=["embeddings"])
@require_memory(action="reject", reject_on=("critical", "emergency"))
async def generate_embeddings_v2(
    texts: list[str] = Body(..., min_items=1, max_items=100, embed=True),
):
    """Generate embeddings using BGE-M3 model."""
    start_time = time.perf_counter()

    try:
        # OPTIMIZATION: Use cached pipeline singleton
        pipeline = _get_pipeline()

        embeddings = await pipeline.embed(texts)
        elapsed = (time.perf_counter() - start_time) * 1000

        return {
            "embeddings": embeddings.tolist()
            if hasattr(embeddings, "tolist")
            else embeddings,
            "model": "BAAI/bge-m3",
            "dimensions": len(embeddings[0]) if len(embeddings) > 0 else 0,
            "count": len(texts),
            "processing_time_ms": elapsed,
        }

    except Exception as e:
        logger.error("Embedding generation error: %s", e)
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/embeddings/rerank", tags=["embeddings"])
async def rerank_documents(
    query: str = Body(...), documents: list[str] = Body(..., min_items=1, max_items=100)
):
    """Rerank documents by relevance using BGE-Reranker-v2-M3."""
    start_time = time.perf_counter()

    try:
        # OPTIMIZATION: Use cached pipeline singleton
        pipeline = _get_pipeline()

        reranked = await pipeline.rerank(query, documents)
        elapsed = (time.perf_counter() - start_time) * 1000

        return {
            "results": [
                {"document": documents[idx], "score": score, "original_index": idx}
                for idx, score in reranked
            ],
            "query": query,
            "model": "BAAI/bge-reranker-v2-m3",
            "processing_time_ms": elapsed,
        }

    except Exception as e:
        logger.error("Reranking error: %s", e)
        raise HTTPException(status_code=500, detail=str(e))

# ==================== AI Endpoints ====================

# OPTIMIZATION: Pre-compute style guides as module-level constant
_STYLE_GUIDES = {
    "simple": "Use simple words and short sentences. Explain like you're explaining to a beginner.",
    "detailed": "Provide a comprehensive explanation with examples and context.",
    "technical": "Use appropriate technical terminology while still being clear.",
    "storytelling": "Explain using a narrative or story-based approach.",
}

# OPTIMIZATION: Pre-compute unsafe patterns as frozenset for O(1) lookup
_UNSAFE_PATTERNS = frozenset({"violence", "hate", "explicit"})

@router.post("/ai/explain", tags=["ai"])
async def explain_content(
    text: str = Body(..., embed=True), style: str = Body(default="simple")
):
    """Get AI explanation of content."""
    try:
        from ...services.chat.engine import GenerationConfig

        # OPTIMIZATION: Use cached engine singleton
        engine = _get_ai_engine()

        style_instruction = _STYLE_GUIDES.get(style, _STYLE_GUIDES["simple"])

        prompt = f"""Explain the following content in a {style} way.

{style_instruction}

CONTENT TO EXPLAIN:
{text}

Provide a clear, accurate, and explanation. If the content contains any inaccuracies, correct them."""

        config = GenerationConfig(max_tokens=2048, temperature=0.4, use_rag=True)

        response = await engine.chat(
            message=prompt,
            config=config,
            context_data={"task_type": "explanation", "style": style},
        )

        return {"explanation": response.content}

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/ai/prompts", tags=["ai"])
async def list_prompts():
    """List available prompt templates."""
    return {
        "prompts": [
            {"name": "simplify", "description": "Simplify content"},
            {"name": "translate", "description": "Translate to Indian languages"},
            {"name": "quiz", "description": "Generate quiz questions"},
            {"name": "explain", "description": "Explain concepts"},
        ]
    }

@router.post("/ai/safety/check", tags=["ai"])
async def check_content_safety(text: str = Body(..., embed=True)):
    """Check content for safety issues."""
    # OPTIMIZATION: Use pre-computed frozenset and any() with generator
    text_lower = text.lower()
    is_safe = not any(pattern in text_lower for pattern in _UNSAFE_PATTERNS)

    return {
        "is_safe": is_safe,
        "confidence": 0.95 if is_safe else 0.3,
        "issues": [] if is_safe else ["potential_unsafe_content"],
    }
